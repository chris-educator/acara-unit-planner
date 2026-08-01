"""Export term unit packs to DOCX, PDF, TXT, and ZIP."""

from __future__ import annotations

import io
import zipfile
from datetime import datetime, timezone

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from fpdf import FPDF
from fpdf.enums import XPos, YPos

_MUTED = RGBColor(0x5C, 0x5C, 0x66)
_PDF_CREATOR = "AppStax ACARA Unit Planner"


def _generated_line() -> str:
    return datetime.now(timezone.utc).strftime("Generated %d %b %Y %H:%M UTC")


def _slug(text: str) -> str:
    slug = "".join(ch if ch.isalnum() or ch in "-_" else "-" for ch in text)[:50]
    return slug.strip("-") or "term-plan"


def _safe_pdf_text(text: str) -> str:
    return (
        (text or "")
        .replace("\u2014", "-")
        .replace("\u2013", "-")
        .replace("\u2019", "'")
        .replace("\u2018", "'")
        .replace("\u201c", '"')
        .replace("\u201d", '"')
        .replace("•", "-")
        .encode("latin-1", "replace")
        .decode("latin-1")
    )


def _header_lines(unit: dict, *, school_name: str = "") -> list[str]:
    title = unit.get("unit_title", "Term Plan")
    lines = [title, "=" * len(title), ""]
    if school_name.strip():
        lines.append(school_name.strip())
    lines.append(f"Topic: {unit.get('topic', '')}")
    lines.append(f"Year level: {unit.get('year_level', '')}")
    lines.append(f"Subject: {unit.get('subject', '')}")
    lines.append(f"Weeks: {unit.get('lesson_count', len(unit.get('lessons') or []))}")
    lines.append(_generated_line())
    return lines


def _lesson_block(lesson: dict) -> list[str]:
    lines = [
        "",
        f"Week {lesson.get('lesson_number')}: {lesson.get('title', '')}",
        "-" * 40,
    ]
    if lesson.get("timing_notes"):
        lines.append(f"Timing: {lesson.get('timing_notes')}")
        lines.append("")
    lines.append("Learning objectives")
    for objective in lesson.get("learning_objectives") or []:
        lines.append(f"• {objective}")
    lines.extend(["", "Materials needed"])
    for material in lesson.get("materials_needed") or []:
        lines.append(f"• {material}")
    lines.extend(
        [
            "",
            "Starter",
            lesson.get("starter", ""),
            "",
            "Main activity",
            lesson.get("main_activity", ""),
            "",
            "Exit ticket",
            lesson.get("exit_ticket", ""),
            "",
            "Differentiation — support",
            lesson.get("differentiation_support", ""),
            "",
            "Differentiation — extension",
            lesson.get("differentiation_extension", ""),
        ]
    )
    return lines


def build_unit_txt(unit: dict, *, school_name: str = "") -> bytes:
    lines = _header_lines(unit, school_name=school_name)
    lines.extend(["", "Unit overview", "-" * 14, unit.get("overview", ""), ""])

    criteria = unit.get("success_criteria") or []
    if criteria:
        lines.extend(["Success criteria", "-" * 16])
        for item in criteria:
            lines.append(f"• {item}")
        lines.append("")

    descriptors = unit.get("suggested_descriptors") or []
    if descriptors:
        lines.extend(["Suggested curriculum links", "-" * 26])
        for item in descriptors:
            lines.append(f"• {item.get('label', '')}: {item.get('summary', '')}")
        lines.append("")

    for lesson in unit.get("lessons") or []:
        lines.extend(_lesson_block(lesson))

    assessment = unit.get("unit_assessment") or {}
    lines.extend(
        [
            "",
            assessment.get("title", "Unit assessment"),
            "-" * len(assessment.get("title", "Unit assessment")),
            assessment.get("instructions", ""),
            "",
        ]
    )
    for index, task in enumerate(assessment.get("tasks") or [], start=1):
        lines.append(f"{index}. {task}")

    rubric = assessment.get("rubric") or []
    if rubric:
        lines.extend(["", "Assessment rubric", "-" * 18])
        for row in rubric:
            lines.extend(
                [
                    "",
                    row.get("criterion", ""),
                    f"  Developing: {row.get('developing', '')}",
                    f"  Meeting: {row.get('meeting', '')}",
                    f"  Exceeding: {row.get('exceeding', '')}",
                ]
            )

    return ("\n".join(lines).strip() + "\n").encode("utf-8")


def _add_rubric_table(doc: Document, rubric: list[dict]) -> None:
    if not rubric:
        return
    doc.add_heading("Assessment rubric", level=2)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = table.rows[0].cells
    for index, label in enumerate(["Criterion", "Developing", "Meeting", "Exceeding"]):
        headers[index].text = label
    for row in rubric:
        cells = table.add_row().cells
        cells[0].text = row.get("criterion", "")
        cells[1].text = row.get("developing", "")
        cells[2].text = row.get("meeting", "")
        cells[3].text = row.get("exceeding", "")


def build_unit_docx(unit: dict, *, school_name: str = "") -> bytes:
    doc = Document()
    title = unit.get("unit_title", "Term Plan")
    heading = doc.add_heading(title, level=0)
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph()
    if school_name.strip():
        meta.add_run(school_name.strip()).bold = True
        meta.add_run("\n")
    meta.add_run(f"Topic: {unit.get('topic', '')}\n")
    meta.add_run(
        f"Year level: {unit.get('year_level', '')} · Subject: {unit.get('subject', '')}\n"
    )
    meta.add_run(_generated_line()).font.color.rgb = _MUTED

    doc.add_heading("Unit overview", level=1)
    doc.add_paragraph(unit.get("overview", ""))

    criteria = unit.get("success_criteria") or []
    if criteria:
        doc.add_heading("Success criteria", level=1)
        for item in criteria:
            doc.add_paragraph(item, style="List Bullet")

    descriptors = unit.get("suggested_descriptors") or []
    if descriptors:
        doc.add_heading("Suggested curriculum links", level=1)
        for item in descriptors:
            doc.add_paragraph(
                f"{item.get('label', '')} — {item.get('summary', '')}",
                style="List Bullet",
            )

    for lesson in unit.get("lessons") or []:
        doc.add_heading(
            f"Week {lesson.get('lesson_number')}: {lesson.get('title', '')}",
            level=1,
        )
        if lesson.get("timing_notes"):
            doc.add_paragraph(lesson.get("timing_notes", "")).italic = True
        doc.add_heading("Learning objectives", level=2)
        for objective in lesson.get("learning_objectives") or []:
            doc.add_paragraph(objective, style="List Bullet")
        doc.add_heading("Materials needed", level=2)
        for material in lesson.get("materials_needed") or []:
            doc.add_paragraph(material, style="List Bullet")
        doc.add_heading("Starter", level=2)
        doc.add_paragraph(lesson.get("starter", ""))
        doc.add_heading("Main activity", level=2)
        doc.add_paragraph(lesson.get("main_activity", ""))
        doc.add_heading("Exit ticket", level=2)
        doc.add_paragraph(lesson.get("exit_ticket", ""))
        doc.add_heading("Differentiation — support", level=2)
        doc.add_paragraph(lesson.get("differentiation_support", ""))
        doc.add_heading("Differentiation — extension", level=2)
        doc.add_paragraph(lesson.get("differentiation_extension", ""))

    assessment = unit.get("unit_assessment") or {}
    doc.add_heading(assessment.get("title", "Unit assessment"), level=1)
    doc.add_paragraph(assessment.get("instructions", ""))
    for index, task in enumerate(assessment.get("tasks") or [], start=1):
        doc.add_paragraph(f"{index}. {task}")
    _add_rubric_table(doc, assessment.get("rubric") or [])

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.font.size is None:
                run.font.size = Pt(11)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _pdf_heading(pdf: FPDF, text: str, *, size: int = 13) -> None:
    pdf.ln(3)
    pdf.set_x(pdf.l_margin)
    pdf.set_font("Helvetica", "B", size)
    pdf.multi_cell(0, 7, _safe_pdf_text(text), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_font("Helvetica", "", 10)


def _pdf_body(pdf: FPDF, text: str) -> None:
    if not (text or "").strip():
        return
    pdf.set_x(pdf.l_margin)
    pdf.multi_cell(0, 5, _safe_pdf_text(text), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(1)


def _pdf_bullets(pdf: FPDF, items: list[str]) -> None:
    for item in items:
        pdf.set_x(pdf.l_margin)
        pdf.multi_cell(0, 5, _safe_pdf_text(f"- {item}"), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(1)


def build_unit_pdf(unit: dict, *, school_name: str = "") -> bytes:
    title = unit.get("unit_title", "Term Plan")
    pdf = FPDF()
    pdf.set_title(_safe_pdf_text(title))
    pdf.set_author("AppStax")
    pdf.set_creator(_PDF_CREATOR)
    pdf.set_subject("ACARA term unit plan")
    if hasattr(pdf, "set_lang"):
        pdf.set_lang("en-AU")
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    pdf.set_font("Helvetica", "B", 16)
    pdf.multi_cell(0, 8, _safe_pdf_text(title), align="C", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(2)
    pdf.set_font("Helvetica", "", 10)
    if school_name.strip():
        pdf.cell(
            0,
            5,
            _safe_pdf_text(school_name.strip()),
            new_x=XPos.LMARGIN,
            new_y=YPos.NEXT,
            align="C",
        )
    pdf.cell(
        0,
        5,
        _safe_pdf_text(f"Topic: {unit.get('topic', '')}"),
        new_x=XPos.LMARGIN,
        new_y=YPos.NEXT,
        align="C",
    )
    pdf.cell(
        0,
        5,
        _safe_pdf_text(
            f"Year level: {unit.get('year_level', '')}  |  Subject: {unit.get('subject', '')}"
        ),
        new_x=XPos.LMARGIN,
        new_y=YPos.NEXT,
        align="C",
    )
    pdf.cell(
        0,
        5,
        _safe_pdf_text(_generated_line()),
        new_x=XPos.LMARGIN,
        new_y=YPos.NEXT,
        align="C",
    )
    pdf.ln(4)

    _pdf_heading(pdf, "Unit overview")
    _pdf_body(pdf, unit.get("overview", ""))

    criteria = unit.get("success_criteria") or []
    if criteria:
        _pdf_heading(pdf, "Success criteria")
        _pdf_bullets(pdf, [str(item) for item in criteria])

    descriptors = unit.get("suggested_descriptors") or []
    if descriptors:
        _pdf_heading(pdf, "Suggested curriculum links")
        _pdf_bullets(
            pdf,
            [f"{item.get('label', '')} - {item.get('summary', '')}" for item in descriptors],
        )

    for lesson in unit.get("lessons") or []:
        _pdf_heading(
            pdf,
            f"Week {lesson.get('lesson_number')}: {lesson.get('title', '')}",
            size=12,
        )
        if lesson.get("timing_notes"):
            pdf.set_font("Helvetica", "I", 10)
            _pdf_body(pdf, lesson.get("timing_notes", ""))
            pdf.set_font("Helvetica", "", 10)
        _pdf_heading(pdf, "Learning objectives", size=11)
        _pdf_bullets(pdf, [str(o) for o in (lesson.get("learning_objectives") or [])])
        _pdf_heading(pdf, "Materials needed", size=11)
        _pdf_bullets(pdf, [str(m) for m in (lesson.get("materials_needed") or [])])
        for label, key in (
            ("Starter", "starter"),
            ("Main activity", "main_activity"),
            ("Exit ticket", "exit_ticket"),
            ("Differentiation - support", "differentiation_support"),
            ("Differentiation - extension", "differentiation_extension"),
        ):
            _pdf_heading(pdf, label, size=11)
            _pdf_body(pdf, lesson.get(key, ""))

    assessment = unit.get("unit_assessment") or {}
    _pdf_heading(pdf, assessment.get("title", "Unit assessment"))
    _pdf_body(pdf, assessment.get("instructions", ""))
    for index, task in enumerate(assessment.get("tasks") or [], start=1):
        _pdf_body(pdf, f"{index}. {task}")

    rubric = assessment.get("rubric") or []
    if rubric:
        _pdf_heading(pdf, "Assessment rubric", size=12)
        for row in rubric:
            pdf.set_font("Helvetica", "B", 10)
            _pdf_body(pdf, row.get("criterion", ""))
            pdf.set_font("Helvetica", "", 10)
            _pdf_body(pdf, f"Developing: {row.get('developing', '')}")
            _pdf_body(pdf, f"Meeting: {row.get('meeting', '')}")
            _pdf_body(pdf, f"Exceeding: {row.get('exceeding', '')}")
            pdf.ln(1)

    out = pdf.output()
    if isinstance(out, (bytes, bytearray)):
        return bytes(out)
    if isinstance(out, str):
        return out.encode("latin-1")
    raise TypeError(f"Unexpected fpdf output type: {type(out)!r}")


def build_export_zip(unit: dict, *, school_name: str = "") -> bytes:
    slug = _slug(unit.get("unit_title", unit.get("topic", "term-plan")))
    docx_bytes = build_unit_docx(unit, school_name=school_name)
    pdf_bytes = build_unit_pdf(unit, school_name=school_name)
    txt_bytes = build_unit_txt(unit, school_name=school_name)

    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w", zipfile.ZIP_DEFLATED) as archive:
        archive.writestr(f"{slug}-term-plan.docx", docx_bytes)
        archive.writestr(f"{slug}-term-plan.pdf", pdf_bytes)
        archive.writestr(f"{slug}-term-plan.txt", txt_bytes)
    return buffer.getvalue()
